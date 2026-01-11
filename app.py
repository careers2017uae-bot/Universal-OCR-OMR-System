# app.py
"""
Universal OCR + OMR System with Document Structure Analysis
Cloud-Native | Donut Model (OCR-Free LayoutLM) | Production-Safe
"""

import streamlit as st
from PIL import Image
import torch
from transformers import DonutProcessor, VisionEncoderDecoderModel
import json
from docx import Document
from fpdf import FPDF
import tempfile
import os
import numpy as np
import cv2

# ---------------- CONFIG ----------------

st.set_page_config(
    page_title="Universal OCR & OMR System",
    page_icon="📄",
    layout="centered"
)

# Initialize Donut model for document understanding
@st.cache_resource
def load_donut_model():
    """Load Donut model for OCR-free document understanding"""
    processor = DonutProcessor.from_pretrained("naver-clova-ix/donut-base-finetuned-cord-v2")
    model = VisionEncoderDecoderModel.from_pretrained("naver-clova-ix/donut-base-finetuned-cord-v2")
    
    if torch.cuda.is_available():
        model = model.to("cuda")
    
    return processor, model

# ---------------- UI ----------------

st.title("📄 Universal OCR & OMR System with Layout Analysis")
st.caption("Document Structure • OMR Detection • Multi-Format Output")

uploaded_file = st.file_uploader(
    "📷 Upload scanned document or form",
    type=["png", "jpg", "jpeg", "pdf"]
)

# Advanced options
col1, col2 = st.columns(2)
with col1:
    task_type = st.selectbox(
        "Task Type",
        ["Document OCR with Layout", "OMR Form Processing", "Handwriting Recognition"]
    )
with col2:
    output_format = st.selectbox(
        "📤 Output Format",
        ["Structured JSON", "Word (.docx)", "PDF (.pdf)", "Markdown"]
    )

process_btn = st.button("🚀 Process Document", type="primary")

# ---------------- CORE FUNCTIONS ----------------

def process_with_donut(image: Image.Image, task: str = "document parsing") -> dict:
    """
    Process document using Donut model for OCR-free document understanding
    Returns structured data with layout information
    """
    try:
        processor, model = load_donut_model()
        
        # Prepare image for Donut model
        pixel_values = processor(image.convert("RGB"), return_tensors="pt").pixel_values
        
        if torch.cuda.is_available():
            pixel_values = pixel_values.to("cuda")
        
        # Generate document understanding output
        task_prompt = f"<s_{task}>"
        decoder_input_ids = processor.tokenizer(
            task_prompt, 
            add_special_tokens=False, 
            return_tensors="pt"
        ).input_ids
        
        # Generate structured output
        outputs = model.generate(
            pixel_values,
            decoder_input_ids=decoder_input_ids,
            max_length=model.decoder.config.max_position_embeddings,
            pad_token_id=processor.tokenizer.pad_token_id,
            eos_token_id=processor.tokenizer.eos_token_id,
            use_cache=True,
            bad_words_ids=[[processor.tokenizer.unk_token_id]],
            return_dict_in_generate=True,
        )
        
        # Decode the output
        sequence = processor.batch_decode(outputs.sequences)[0]
        sequence = sequence.replace(processor.tokenizer.eos_token, "").replace(processor.tokenizer.pad_token, "")
        sequence = sequence.replace("<s_"+task+">", "").replace("</s_"+task+">", "")
        
        # Parse JSON output
        try:
            result = json.loads(sequence)
        except json.JSONDecodeError:
            result = {"text": sequence, "layout": "unstructured"}
        
        return result
        
    except Exception as e:
        st.error(f"Donut model error: {str(e)}")
        return {"error": str(e), "text": "", "layout": {}}

def detect_omr_markers(image: Image.Image) -> dict:
    """
    Detect OMR (bubble/checkbox) markers using computer vision
    """
    # Convert PIL to OpenCV format
    opencv_image = cv2.cvtColor(np.array(image.convert("RGB")), cv2.COLOR_RGB2BGR)
    gray = cv2.cvtColor(opencv_image, cv2.COLOR_BGR2GRAY)
    
    # Apply adaptive thresholding
    _, thresh = cv2.threshold(gray, 150, 255, cv2.THRESH_BINARY_INV)
    
    # Find contours (potential bubbles)
    contours, _ = cv2.findContours(thresh, cv2.RETR_EXTERNAL, cv2.CHAIN_APPROX_SIMPLE)
    
    omr_results = {
        "total_bubbles_detected": len(contours),
        "filled_bubbles": 0,
        "bubble_details": []
    }
    
    min_area = 100  # Minimum area to be considered a bubble
    for i, contour in enumerate(contours):
        area = cv2.contourArea(contour)
        if area > min_area:
            # Calculate filled ratio
            x, y, w, h = cv2.boundingRect(contour)
            aspect_ratio = w / h if h != 0 else 0
            
            # Check if it's circular/elliptical (bubble-like)
            if 0.7 < aspect_ratio < 1.3:
                # Calculate filled percentage
                bubble_region = thresh[y:y+h, x:x+w]
                filled_pixels = np.sum(bubble_region > 0)
                total_pixels = w * h
                filled_ratio = filled_pixels / total_pixels if total_pixels > 0 else 0
                
                is_filled = filled_ratio > 0.3  # Threshold for filled bubble
                
                omr_results["bubble_details"].append({
                    "id": i,
                    "position": {"x": x, "y": y, "width": w, "height": h},
                    "area": area,
                    "filled_ratio": filled_ratio,
                    "is_filled": is_filled
                })
                
                if is_filled:
                    omr_results["filled_bubbles"] += 1
    
    return omr_results

def generate_structured_output(content: dict, format_type: str) -> str:
    """Generate output in specified format with document structure preserved"""
    
    if format_type == "Structured JSON":
        # Save JSON to temporary file
        path = tempfile.mktemp(suffix=".json")
        with open(path, "w") as f:
            json.dump(content, f, indent=2)
        return path
    
    elif format_type == "Word (.docx)":
        doc = Document()
        
        # Add title if available
        if "header" in content:
            doc.add_heading(content.get("header", "Document"), level=1)
        
        # Add structured content
        if "sections" in content:
            for section in content.get("sections", []):
                doc.add_heading(section.get("title", "Section"), level=2)
                for para in section.get("paragraphs", []):
                    doc.add_paragraph(para)
        
        # Add OMR results if available
        if "omr_results" in content:
            doc.add_heading("OMR Results", level=2)
            omr = content["omr_results"]
            doc.add_paragraph(f"Total bubbles detected: {omr.get('total_bubbles_detected', 0)}")
            doc.add_paragraph(f"Filled bubbles: {omr.get('filled_bubbles', 0)}")
        
        path = tempfile.mktemp(suffix=".docx")
        doc.save(path)
        return path
    
    elif format_type == "PDF (.pdf)":
        pdf = FPDF()
        pdf.set_auto_page_break(auto=True, margin=15)
        pdf.add_page()
        pdf.set_font("Arial", size=11)
        
        # Add content
        if "text" in content:
            for line in content["text"].split("\n"):
                if line.strip():
                    pdf.multi_cell(0, 8, line)
        
        path = tempfile.mktemp(suffix=".pdf")
        pdf.output(path)
        return path
    
    else:  # Markdown
        md_content = ""
        if "header" in content:
            md_content += f"# {content['header']}\n\n"
        
        if "sections" in content:
            for section in content["sections"]:
                md_content += f"## {section.get('title', 'Section')}\n\n"
                for para in section.get("paragraphs", []):
                    md_content += f"{para}\n\n"
        
        path = tempfile.mktemp(suffix=".md")
        with open(path, "w") as f:
            f.write(md_content)
        return path

# ---------------- MAIN PIPELINE ----------------

if process_btn and uploaded_file:
    try:
        # Load and process image
        image = Image.open(uploaded_file)
        
        with st.spinner("🔍 Analyzing document structure with Donut model..."):
            # Process based on task type
            if task_type == "OMR Form Processing":
                # OMR processing
                omr_results = detect_omr_markers(image)
                
                # Also extract any text present
                doc_results = process_with_donut(image, "document parsing")
                
                combined_results = {
                    "document_type": "OMR Form",
                    "omr_results": omr_results,
                    "extracted_text": doc_results.get("text", ""),
                    "layout": doc_results.get("layout", {})
                }
                
                # Display OMR visualization
                if omr_results["bubble_details"]:
                    st.subheader("OMR Analysis Results")
                    col1, col2 = st.columns(2)
                    with col1:
                        st.metric("Total Bubbles", omr_results["total_bubbles_detected"])
                    with col2:
                        st.metric("Filled Bubbles", omr_results["filled_bubbles"])
                    
                    # Show sample bubbles
                    if len(omr_results["bubble_details"]) > 0:
                        st.write("Sample detected bubbles:")
                        st.json(omr_results["bubble_details"][:3])  # Show first 3
                
                content = combined_results
                
            else:
                # Standard document processing
                doc_results = process_with_donut(
                    image, 
                    "handwritten text recognition" if "Handwriting" in task_type 
                    else "document parsing"
                )
                
                # Structure the output
                structured_content = {
                    "document_type": task_type,
                    "extraction_timestamp": st.session_state.get("extraction_time", ""),
                    "text": doc_results.get("text", ""),
                    "structured_data": doc_results,
                    "confidence_score": 0.95  # Placeholder for actual confidence
                }
                
                content = structured_content
            
            # Generate output file
            output_path = generate_structured_output(content, output_format)
            
            st.success("✅ Document processing completed!")
            
            # Show preview for text-based formats
            if output_format != "Structured JSON" and "text" in content:
                with st.expander("📝 Extracted Text Preview"):
                    st.text(content["text"][:500] + "..." if len(content["text"]) > 500 else content["text"])
            
            # Download button
            with open(output_path, "rb") as f:
                file_ext = output_format.split(" ")[-1].strip("()")
                st.download_button(
                    f"⬇ Download {output_format}",
                    f,
                    file_name=f"processed_document{file_ext}",
                    mime=f"application/{file_ext.replace('.', '')}"
                )
    
    except Exception as e:
        st.error(f"❌ Processing failed: {str(e)}")
        st.info("Try with a clearer image or different document type")

# ---------------- SIDEBAR SETTINGS ----------------

with st.sidebar:
    st.header("⚙️ Advanced Settings")
    
    confidence_threshold = st.slider(
        "Confidence Threshold", 
        min_value=0.5, 
        max_value=1.0, 
        value=0.7
    )
    
    language = st.selectbox(
        "Primary Language",
        ["Auto-detect", "English", "Spanish", "French", "German", "Chinese", "Japanese"]
    )
    
    preserve_layout = st.checkbox("Preserve Document Layout", value=True)
    
    if st.button("Clear Cache"):
        st.cache_resource.clear()
        st.success("Cache cleared!")

st.caption("Powered by Donut Model • Document Structure Understanding • OMR Detection")
